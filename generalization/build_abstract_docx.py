"""Generate the submittable Word document for the generalization short paper.

Content mirrors generalization/ABSTRACT.md + EXPLAINER.md (honest framing:
label-free method matches supervised selection and is more stable than variance;
cross-domain transfer; ME-HSI noted slightly; NO "beats random" / "first" claims).

Run: python generalization/build_abstract_docx.py
Output: generalization/Meloyan_GeneralChannelSelection.docx
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
FIGS = HERE / "figures"
OUT = HERE / "Meloyan_GeneralChannelSelection.docx"

BODY_FONT = "Times New Roman"
BODY_PT = 9.5


def set_columns(section, num, space_twips=360):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols"); sectPr.append(cols)
    cols.set(qn("w:num"), str(num)); cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")


def tighten(p, before=0, after=2, line=1.0):
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line


def _margins(section):
    section.top_margin = Inches(0.75); section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.625); section.right_margin = Inches(0.625)


def body_para(doc, text, *, justify=True, first_indent=0.18, size=BODY_PT,
              bold=False, italic=False):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    tighten(p)
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(first_indent)
    r = p.add_run(text)
    r.font.name = BODY_FONT; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p


def section_heading(doc, numeral, title):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(p, before=6, after=2)
    r = p.add_run(f"{numeral}. {title.upper()}")
    r.font.name = BODY_FONT; r.font.size = Pt(BODY_PT); r.bold = True; r.font.small_caps = True


def _caption(doc, text):
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(cap, before=1, after=4)
    r = cap.add_run(text); r.font.name = BODY_FONT; r.font.size = Pt(8)


def column_figure(doc, img, caption, width_in=3.35):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(p, before=3, after=0)
    if Path(img).exists():
        p.add_run().add_picture(str(img), width=Inches(width_in))
    else:
        p.add_run(f"[missing figure: {Path(img).name}]")
    _caption(doc, caption)


def fullwidth_figure(doc, img, caption, width_in=6.9):
    s1 = doc.add_section(WD_SECTION.CONTINUOUS); _margins(s1); set_columns(s1, 1)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(p, before=2, after=0)
    if Path(img).exists():
        p.add_run().add_picture(str(img), width=Inches(width_in))
    else:
        p.add_run(f"[missing figure: {Path(img).name}]")
    _caption(doc, caption)
    s2 = doc.add_section(WD_SECTION.CONTINUOUS); _margins(s2); set_columns(s2, 2)


ABSTRACT = (
    "Modern sensing systems produce many channels that are highly correlated and organized "
    "into coupled groups. Choosing a small, interpretable subset of the original channels—"
    "without labels—reduces cost and complexity, but existing options force a trade-off: "
    "unsupervised filters (variance, PCA) ignore inter-channel dependency, while "
    "dependency-aware selectors (mRMR, JMI) require labels. We present an unsupervised method "
    "that captures channel dependency directly. A group-structured autoencoder learns a joint "
    "representation across channel groups; each channel's importance is measured by the "
    "sensitivity of the reconstruction to perturbations of the learned latent factors; and a "
    "relevance-redundancy criterion selects a diverse, informative subset of the actual "
    "channels. The method is discrete—it returns real channels, not a projection—and "
    "interpretable. Originally developed for multi-excitation hyperspectral imaging, where it "
    "identifies discriminative wavelength bands and outperforms classical band-selection "
    "baselines, the same selection procedure transfers to a very different modality with only "
    "a domain-appropriate encoder. On a standard wearable-sensor human-activity-recognition "
    "benchmark, evaluated with leave-one-subject-out cross-validation, the label-free method "
    "matches supervised mutual-information selection and is substantially more stable across "
    "subjects than variance-based selection. A single dependency-aware, label-free selection "
    "principle thus generalizes across markedly different sensor modalities."
)

REFS = [
    "M. F. Balin, A. Abid, and J. Zou, \"Concrete autoencoders: Differentiable feature "
    "selection and reconstruction,\" in Proc. ICML, 2019.",
    "H. Peng, F. Long, and C. Ding, \"Feature selection based on mutual information: "
    "Criteria of max-dependency, max-relevance, and min-redundancy,\" IEEE Trans. Pattern "
    "Anal. Mach. Intell., vol. 27, no. 8, 2005.",
    "X. He, D. Cai, and P. Niyogi, \"Laplacian score for feature selection,\" in Proc. "
    "NeurIPS, 2005.",
    "Y. Cai et al., \"BS-Nets: An end-to-end framework for band selection of hyperspectral "
    "image,\" IEEE Trans. Geosci. Remote Sens., vol. 58, no. 3, 2020.",
    "A. Reiss and D. Stricker, \"Introducing a new benchmarked dataset for activity "
    "monitoring,\" in Proc. IEEE Int. Symp. Wearable Computers (ISWC), 2012.",
    "N. Meloyan and N. Sarvazyan, \"Perturbation-based wavelength selection for "
    "multi-excitation hyperspectral imaging\" (companion work).",
]


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Inches(11); sec.page_width = Inches(8.5); _margins(sec)

    # Title
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER; tighten(t, after=4)
    rt = t.add_run("Label-Free, Dependency-Aware Channel Selection\n"
                   "for Coupled Multi-Channel Sensor Data")
    rt.font.name = BODY_FONT; rt.font.size = Pt(18); rt.bold = True

    a = doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER; tighten(a, before=2, after=0)
    ra = a.add_run("Narek Meloyan, Aleksandr Hayrapetyan, and Narine Sarvazyan")
    ra.font.name = BODY_FONT; ra.font.size = Pt(11)
    af = doc.add_paragraph(); af.alignment = WD_ALIGN_PARAGRAPH.CENTER; tighten(af, before=1, after=8)
    raf = af.add_run("American University of Armenia, Yerevan, Armenia\n"
                     "{narek.meloyan, aleksandr.hayrapetyan, nsarvazyan}@aua.am")
    raf.font.name = BODY_FONT; raf.font.size = Pt(9); raf.italic = True

    body = doc.add_section(WD_SECTION.CONTINUOUS); _margins(body); set_columns(body, 2)

    # Abstract
    ab = doc.add_paragraph(); ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; tighten(ab, after=3)
    lab = ab.add_run("Abstract—"); lab.font.name = BODY_FONT; lab.font.size = Pt(BODY_PT)
    lab.bold = True; lab.italic = True
    rab = ab.add_run(ABSTRACT); rab.font.name = BODY_FONT; rab.font.size = Pt(BODY_PT)
    rab.bold = True; rab.italic = True

    idx = doc.add_paragraph(); idx.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; tighten(idx, after=4)
    il = idx.add_run("Index Terms—"); il.font.name = BODY_FONT; il.font.size = Pt(BODY_PT)
    il.bold = True; il.italic = True
    it = idx.add_run("unsupervised feature selection, channel selection, conditional relevance, "
                     "autoencoders, interpretability, human activity recognition, "
                     "hyperspectral imaging.")
    it.font.name = BODY_FONT; it.font.size = Pt(BODY_PT); it.italic = True

    # I. Introduction
    section_heading(doc, "I", "Introduction")
    body_para(doc,
        "Complex sensors emit many channels that overlap heavily and arrive in natural groups—"
        "the three axes of one inertial unit, or the emission bands of one excitation. For cost, "
        "power, and interpretability one often wants to retain only a handful of the actual "
        "channels, not a learned mixture of them. Existing selectors force a trade-off. "
        "Unsupervised filters such as variance ranking and PCA judge each channel by its marginal "
        "statistics and ignore inter-channel dependency. Dependency-aware selectors such as mRMR "
        "and JMI model conditional relevance but require labels. The cell that is "
        "label-free, dependency-aware, and discrete is sparsely populated; this work targets it "
        "and, more importantly, shows that one such method generalizes across very different "
        "sensor modalities.")

    # II. Method
    section_heading(doc, "II", "Method")
    body_para(doc,
        "The method has three stages (Fig. 1). First, a group-structured autoencoder is trained "
        "by reconstruction alone—no labels: each channel group is encoded separately, the "
        "per-group features are mean-fused into a shared latent representation, and the latent is "
        "decoded back to each group. Second, perturbation-based attribution nudges individual "
        "latent factors and measures how much each channel's reconstruction changes; channels the "
        "model relies on react strongly. Third, a relevance-redundancy (maximal marginal "
        "relevance) criterion greedily selects channels that are individually informative yet "
        "non-redundant with those already chosen. The output is a discrete, ordered, interpretable "
        "subset of the real channels that can be thresholded to any target size.")
    fullwidth_figure(doc, FIGS / "fig_method.png",
        "Fig. 1.  The selection pipeline. The engine (perturbation → per-channel influence "
        "→ relevance-redundancy selection) is identical across domains; only the "
        "encoder/decoder changes with the data's regular axis.")

    # III. Cross-domain transfer
    section_heading(doc, "III", "Cross-Domain Transfer")
    body_para(doc,
        "The selection engine is modality-agnostic: it operates on an abstract layout of "
        "groups, channels, and one regular axis. Moving to a new modality requires only an "
        "encoder matched to that axis (Fig. 2). The method was originally developed for "
        "multi-excitation hyperspectral imaging, where a group is an excitation wavelength, a "
        "channel is an emission band, the regular axis is 2-D space, and a 3-D convolutional "
        "encoder is used; there it identifies discriminative bands and outperforms classical "
        "band-selection baselines. For wearable human-activity recognition, a group is a "
        "body-worn inertial unit, a channel is one sensor axis, and the regular axis is time, so "
        "a 1-D convolutional encoder is substituted. No part of the selection logic was rewritten.")
    fullwidth_figure(doc, FIGS / "fig_crossdomain.png",
        "Fig. 2.  The same method instantiated in two modalities. Only the encoder changes.")

    # IV. Experiments
    section_heading(doc, "IV", "Experiments and Results")
    body_para(doc,
        "We evaluate on a standard wearable-sensor activity-recognition benchmark with three "
        "body-worn inertial units, using leave-one-subject-out cross-validation and macro-F1, the "
        "rigorous protocol in which the test subject is never seen during training. We compare the "
        "label-free method against a supervised mutual-information selector, an unsupervised "
        "variance baseline, and random selection, sweeping the number of retained channels.")
    column_figure(doc, FIGS / "fig_acc_vs_k.png",
        "Fig. 3.  Macro-F1 vs. number of selected channels (leave-one-subject-out). The "
        "label-free method (red) tracks the supervised selector (blue) and exceeds variance "
        "(orange).")
    body_para(doc,
        "Two results stand out. First, the label-free method matches the supervised "
        "mutual-information selector across channel budgets, achieving without labels what the "
        "supervised method requires labels to attain. Second, it consistently exceeds the "
        "variance baseline and—crucially for deployment—selects far more stably across subjects "
        "(Fig. 4): variance selection swings widely from subject to subject, whereas the proposed "
        "method does not. We note in the interest of full disclosure that on this particular "
        "benchmark no selector, including the supervised one, surpasses random selection: its "
        "channels are intrinsically redundant, so almost any moderate subset approaches the "
        "all-channel ceiling. The comparison against a supervised selector and the stability "
        "analysis are therefore the informative axes here.")
    column_figure(doc, FIGS / "fig_stability.png",
        "Fig. 4.  Stability of the selected set across subjects (lower is better). The proposed "
        "method is markedly more consistent than variance selection.")

    # V. Conclusion
    section_heading(doc, "V", "Conclusion and Future Work")
    body_para(doc,
        "A single unsupervised, dependency-aware selection principle transfers from hyperspectral "
        "imaging to wearable sensors with only an encoder substitution, matching supervised "
        "selection without labels and selecting more stably than variance. Future work targets a "
        "higher-redundancy activity benchmark, where intelligent selection should separate from "
        "random by a clear margin, and a direct comparison against differentiable unsupervised "
        "selectors such as the concrete autoencoder.")

    # Acknowledgement / AI disclosure
    section_heading(doc, "VI", "Acknowledgement")
    body_para(doc,
        "AI-assisted tooling was used for code scaffolding, figure generation, and manuscript "
        "drafting; all experimental design, analysis, and claims were verified by the authors.")

    # References
    section_heading(doc, "", "References")
    for i, r in enumerate(REFS, 1):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        tighten(p, after=1)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        run = p.add_run(f"[{i}] {r}")
        run.font.name = BODY_FONT; run.font.size = Pt(8)

    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    build()
