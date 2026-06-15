"""Generate the CODASSCA 2026 short-paper submission (3 pages, IEEE-style).

Framing (decided 2026-05-23): the paper is GENERALIZED away from the
multi-excitation hyperspectral / perturbation-autoencoder branding used in the
TPAMI journal paper and the IASIM 2026 poster. Here the contribution is framed
as an *unsupervised, dependency-aware dimensionality-reduction* method that
captures CONDITIONAL channel relevance, with multi-excitation fluorescence
imaging used only as a hard BENCHMARK. This makes the submission a distinct
contribution (general method) rather than a re-submission of the ME-HSI work,
and targets CODASSCA Track 1 (Data Science & Information-Theoretic Approaches).

Format target: IEEE two-column, ~3 pages (short-paper category).
NOTE: CODASSCA mandates the official IEEE proceedings template. This DOCX is a
content draft to be poured into that template before submission. It also
includes the required AI-use disclosure in the acknowledgements.

Run:
    .venv/bin/python CODASSCA2026_Submission/build_codassca_shortpaper.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
OUT = HERE / "CODASSCA2026_Meloyan_ShortPaper.docx"

# Figures
FIG_ARCH = HERE / "fig_architecture.png"
FIG_ENVELOPE = ROOT / "MasterThesis_Narek_Meloyan" / "figures" / "accuracy_envelope.png"
FIG_RELMAP = ROOT / "MasterThesis_Narek_Meloyan" / "figures" / "wavelength_heatmap.png"

BODY_FONT = "Times New Roman"
BODY_PT = 9.5


# --------------------------------------------------------------------------
# Low-level helpers
# --------------------------------------------------------------------------
def set_columns(section, num: int, space_twips: int = 360) -> None:
    """Force a section into `num` newspaper-style columns."""
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")


def tighten(paragraph, before=0, after=2, line=1.0) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def body_para(doc, text, *, justify=True, first_indent=0.18,
              size=BODY_PT, bold=False, italic=False):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    tighten(p)
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(first_indent)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def section_heading(doc, numeral, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(p, before=6, after=2)
    r = p.add_run(f"{numeral}. {title.upper()}")
    r.font.name = BODY_FONT
    r.font.size = Pt(BODY_PT)
    r.bold = True
    r.font.small_caps = True
    return p


def sub_heading(doc, letter, title):
    p = doc.add_paragraph()
    tighten(p, before=4, after=1)
    r = p.add_run(f"{letter}. {title}")
    r.font.name = BODY_FONT
    r.font.size = Pt(BODY_PT)
    r.italic = True
    return p


def _set_margins(section):
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)


def _caption(doc, text):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(cap, before=1, after=4)
    r = cap.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(8)
    return cap


def column_figure(doc, img_path, caption, width_in=3.35):
    """Insert a single-column-width figure inline in the 2-column flow."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(p, before=3, after=0)
    if Path(img_path).exists():
        p.add_run().add_picture(str(img_path), width=Inches(width_in))
    else:
        p.add_run(f"[missing figure: {Path(img_path).name}]")
    _caption(doc, caption)


def fullwidth_figure(doc, img_path, caption, width_in=6.9):
    """Span both columns: break to 1 column, place figure, resume 2 columns."""
    s1 = doc.add_section(WD_SECTION.CONTINUOUS)
    _set_margins(s1)
    set_columns(s1, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(p, before=2, after=0)
    if Path(img_path).exists():
        p.add_run().add_picture(str(img_path), width=Inches(width_in))
    else:
        p.add_run(f"[missing figure: {Path(img_path).name}]")
    _caption(doc, caption)
    s2 = doc.add_section(WD_SECTION.CONTINUOUS)
    _set_margins(s2)
    set_columns(s2, 2)


def results_table(doc):
    """Compact reduction-vs-accuracy summary (Table I)."""
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(cap, before=4, after=1)
    rc = cap.add_run("TABLE I.  Accuracy is maintained or improved under "
                     "aggressive channel reduction.")
    rc.font.name = BODY_FONT; rc.font.size = Pt(8); rc.bold = True
    rc.font.small_caps = True

    rows = [
        ("Dataset", "Channels", "Reduction", "Accuracy"),
        ("Biological", "192 (all)", "—", "88.2%"),
        ("Biological", "80", "58%", "95.2%"),
        ("Biological", "9", "95%", "89.4%"),
        ("Materials", "158 (all)", "—", "79.8%"),
        ("Materials", "30", "81%", "85.6%"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    table.alignment = 1
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            par = cell.paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tighten(par, before=0, after=0)
            run = par.add_run(val)
            run.font.name = BODY_FONT
            run.font.size = Pt(8)
            run.bold = (i == 0)
    return table


# --------------------------------------------------------------------------
# Build
# --------------------------------------------------------------------------
def build() -> None:
    doc = Document()

    # Page geometry (US Letter, IEEE-ish margins)
    sec = doc.sections[0]
    sec.page_height = Inches(11)
    sec.page_width = Inches(8.5)
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(0.625)
    sec.right_margin = Inches(0.625)

    # ---- Title block (single column, spans full width) ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(title, before=0, after=4)
    rt = title.add_run("Dependency-Aware Dimensionality Reduction for "
                       "Complex Sensor Data:\nCapturing Conditional Channel Relevance")
    rt.font.name = BODY_FONT
    rt.font.size = Pt(20)
    rt.bold = True

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(authors, before=2, after=0)
    ra = authors.add_run("Narek Meloyan, Aleksandr Hayrapetyan, and Narine Sarvazyan")
    ra.font.name = BODY_FONT
    ra.font.size = Pt(11)

    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(aff, before=1, after=8)
    raf = aff.add_run("American University of Armenia, Yerevan, Armenia\n"
                      "{narek.meloyan, aleksandr.hayrapetyan, nsarvazyan}@aua.am")
    raf.font.name = BODY_FONT
    raf.font.size = Pt(9)
    raf.italic = True

    # ---- Switch to two columns for the body ----
    body = doc.add_section(WD_SECTION.CONTINUOUS)
    body.top_margin = Inches(0.75)
    body.bottom_margin = Inches(1.0)
    body.left_margin = Inches(0.625)
    body.right_margin = Inches(0.625)
    set_columns(body, 2)

    # ---- Abstract ----
    abs = doc.add_paragraph()
    abs_label = abs.add_run("Abstract—")
    abs_label.font.name = BODY_FONT
    abs_label.font.size = Pt(BODY_PT)
    abs_label.bold = True
    abs_label.italic = True
    abstract_text = (
        "Modern sensing systems—hyperspectral cameras, multimodal IoT "
        "arrays, and dense instrument suites—routinely produce hundreds of "
        "correlated channels per sample, most of which are redundant or "
        "noise-dominated. Reducing this dimensionality is essential for "
        "storage, transmission, real-time inference, and economical sensor "
        "design. Yet the dominant techniques—principal component analysis "
        "and filter-based ranking—evaluate channels by their marginal "
        "statistics and therefore overlook conditional relevance: the fact "
        "that a channel’s usefulness depends on which other channels are "
        "already retained. We present an unsupervised framework that captures "
        "this dependency directly. A neural encoder with a cross-group fusion "
        "mechanism learns a compact representation of the joint channel "
        "structure; a sensitivity-based attribution then quantifies each "
        "channel’s conditional contribution to that representation; and a "
        "relevance–redundancy criterion extracts a minimal, non-redundant "
        "subset. The method uses no labels and yields an ordered, interpretable "
        "ranking that can be thresholded to any target size. We benchmark on "
        "multi-excitation fluorescence imaging—a deliberately hard regime "
        "with strongly coupled channel groups—across biological and "
        "materials datasets, reducing channels by 80–95% while matching or "
        "exceeding full-data classification accuracy. A 10,000-trial randomized "
        "test confirms the selected channels are far more informative than "
        "chance (12σ separation). Modeling conditional rather than marginal "
        "relevance yields more efficient and more accurate reduction for "
        "complex sensor data."
    )
    ra2 = abs.add_run(abstract_text)
    ra2.font.name = BODY_FONT
    ra2.font.size = Pt(BODY_PT)
    ra2.bold = True
    ra2.italic = True
    abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    tighten(abs, after=3)

    idx = doc.add_paragraph()
    il = idx.add_run("Index Terms—")
    il.font.name = BODY_FONT; il.font.size = Pt(BODY_PT); il.bold = True; il.italic = True
    it = idx.add_run("dimensionality reduction, channel selection, conditional "
                     "relevance, information-theoretic learning, unsupervised "
                     "representation learning, sensor data analytics, "
                     "interpretability, hyperspectral imaging.")
    it.font.name = BODY_FONT; it.font.size = Pt(BODY_PT); it.italic = True
    idx.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    tighten(idx, after=4)

    # ---- I. Introduction ----
    section_heading(doc, "I", "Introduction")
    body_para(doc,
        "Contemporary sensing platforms generate data of rapidly growing "
        "dimensionality. A single hyperspectral frame, a multimodal IoT node, "
        "or an instrumented industrial process can report hundreds of "
        "measurement channels per sample. These channels are rarely "
        "independent: they are physically coupled, spatially or spectrally "
        "adjacent, and heavily redundant. Acquiring, storing, transmitting, "
        "and analyzing the full channel set is costly, and for real-time or "
        "embedded deployment it is often infeasible. Dimensionality reduction "
        "(DR)—retaining only the channels that carry decision-relevant "
        "information—is therefore a prerequisite for efficient and "
        "economical sensing.", first_indent=0)
    body_para(doc,
        "The two dominant DR families both have structural limitations on such "
        "data. Projection methods such as principal component analysis (PCA) "
        "produce linear mixtures of all channels; they compress storage but do "
        "not reduce the number of channels that must be physically acquired, "
        "and the resulting components are not interpretable as instrument "
        "settings. Filter-based selection methods rank channels by marginal "
        "statistics—variance, entropy, or correlation with a target—"
        "and keep the top scorers. Both implicitly treat channels in "
        "isolation.")
    body_para(doc,
        "Isolation is precisely the wrong assumption for complex sensor data. "
        "The informativeness of a channel is conditional: a channel that is "
        "uninformative on its own can be highly discriminative when combined "
        "with another, while two individually strong channels may be mutually "
        "redundant. In information-theoretic terms, useful selection should "
        "maximize the joint relevance of a subset — closer to "
        "I(X_S; Y) than to the sum of marginal terms Σ I(X_i; Y). "
        "Conditional-mutual-information selectors (mRMR, CMIM, JMI) formalize "
        "this, but they are supervised and rely on density estimation that is "
        "unreliable for high-dimensional, continuous, unlabeled channels.")
    body_para(doc,
        "We address the empty cell in this design space: an unsupervised, "
        "dependency-aware DR method that (i) learns the joint channel structure "
        "with a neural representation, (ii) attributes a conditional relevance "
        "score to each channel through sensitivity analysis of that "
        "representation, and (iii) selects a minimal, non-redundant subset via "
        "a relevance–redundancy criterion. The method requires no labels, "
        "returns a discrete and ordered list of physical channels, and is "
        "interpretable. We evaluate it on multi-excitation fluorescence "
        "imaging — chosen as a benchmark because its channel groups are "
        "strongly coupled, making it a stress test for dependency-aware "
        "reduction.")

    # ---- II. Background ----
    section_heading(doc, "II", "Background and Related Work")
    body_para(doc,
        "Marginal selection. Variance, entropy, and correlation filters score "
        "each channel independently. They are fast and unsupervised but blind "
        "to inter-channel dependency, so they over-select correlated channels "
        "and discard jointly informative ones.", first_indent=0)
    body_para(doc,
        "Information-theoretic selection. mRMR [1], CMIM, and JMI explicitly "
        "trade relevance against redundancy using (conditional) mutual "
        "information. They capture dependency but are supervised and require "
        "MI estimation that degrades in high dimensions and is undefined "
        "without labels.")
    body_para(doc,
        "Projection and embedding. PCA, the minimum-noise fraction, "
        "autoencoders [2], and neighbor embeddings such as UMAP [3] yield "
        "mixtures or low-dimensional coordinates rather than channel subsets, "
        "with no path back to the physical channels an engineer must "
        "instrument.")
    body_para(doc,
        "Deep selection. Attention- and reconstruction-based selectors (e.g., "
        "BS-Net [4]) learn channel importance jointly with a task, but are "
        "typically supervised and assume a fixed input grid. None of these "
        "lines simultaneously offer unsupervised operation, conditional "
        "relevance, discrete channel output, and interpretability — the "
        "combination we target.")

    # ---- III. Method ----
    section_heading(doc, "III", "Method")
    sub_heading(doc, "A", "Joint representation with cross-group fusion")
    body_para(doc,
        "The channels are partitioned into G groups that share an acquisition "
        "condition (in the benchmark, an excitation wavelength). As shown in "
        "Fig. 1, each group is encoded by an independent convolutional branch; "
        "the branch outputs are fused by averaging and refined by a shared "
        "convolution into a latent code z. Averaging in latent "
        "space—rather than concatenation—forces the encoder to find a "
        "representation useful across all groups simultaneously, which is what "
        "encodes cross-group dependency: a concatenated code could keep groups "
        "siloed, whereas the average can be reconstructed well only if it "
        "captures the structure shared across groups. The decoder is "
        "symmetric—a mirrored shared convolution expands z before "
        "per-group convolutional branches reconstruct every group; training "
        "minimizes masked reconstruction error and uses no labels. "
        "Group-wise branches also accommodate a variable number of valid "
        "channels per group without padding, and the model degrades gracefully "
        "if a group is missing at inference time.",
        first_indent=0)
    fullwidth_figure(doc, FIG_ARCH,
        "Fig. 1.  Dependency-aware autoencoder. Each channel group is processed "
        "by a parallel convolutional branch; the branch outputs are fused by "
        "averaging, refined by a shared convolution into the latent code z, and "
        "symmetrically expanded by a mirrored shared convolution before parallel "
        "decoders reconstruct every group. Latent averaging is what forces the "
        "model to capture cross-group (conditional) structure.")
    sub_heading(doc, "B", "Conditional attribution")
    body_para(doc,
        "After training, we probe the learned representation. For each "
        "high-variance latent dimension d we apply a small bipolar "
        "perturbation z ± εσ_d e_d, decode both, and measure the "
        "induced change at every input channel,",
        first_indent=0)
    eq = doc.add_paragraph(); eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(eq, before=1, after=1)
    er = eq.add_run("Δ_d(c) = | G(z + εσ_d e_d) "
                    "− G(z − εσ_d e_d) | / (2εσ_d),")
    er.font.name = BODY_FONT; er.font.size = Pt(BODY_PT); er.italic = True
    body_para(doc,
        "a finite-difference estimate of the decoder’s sensitivity "
        "|∂G/∂z_d| at channel c. Aggregating across dimensions, "
        "weighted by each dimension’s score, gives a per-channel influence "
        "value. Because z encodes the joint structure, this influence reflects "
        "a channel’s contribution in the context of all others — a "
        "conditional, rather than marginal, measure of relevance. Unlike "
        "variance ranking, it can surface low-variance channels that are "
        "decisive only in combination. Fig. 2 visualizes the resulting "
        "conditional-relevance scores across the full channel grid of the "
        "benchmark; the structure is concentrated and physically "
        "interpretable, not spread uniformly.")
    fullwidth_figure(doc, FIG_RELMAP,
        "Fig. 2.  Conditional-relevance map over the benchmark’s channel grid "
        "(here indexed by excitation and emission wavelength). Bright cells are "
        "channels the learned representation depends on most; grey cells are "
        "physically invalid. Relevance concentrates in a few coherent regions, "
        "showing the method isolates a compact, interpretable channel set.")
    sub_heading(doc, "C", "Relevance–redundancy selection")
    body_para(doc,
        "Taking the top channels by influence alone yields tight clusters of "
        "mutually redundant channels. We instead select greedily with a "
        "maximum-marginal-relevance criterion [5]: at each step we add the "
        "channel maximizing λ·Rel(c) − (1−λ)·"
        "max_{c′∈S} Sim(c, c′), where Rel is the influence "
        "score, Sim is the cosine similarity between channel profiles, and "
        "λ=0.5 balances relevance against redundancy. This is an "
        "unsupervised, learned surrogate for the relevance–redundancy "
        "objective of conditional-MI selection, and it returns an ordered list "
        "thresholdable to any target dimensionality.", first_indent=0)

    # ---- IV. Benchmark and Results ----
    section_heading(doc, "IV", "Benchmark and Results")
    body_para(doc,
        "Why this benchmark. Multi-excitation fluorescence imaging records the "
        "same scene under several excitation conditions; each condition yields "
        "an emission spectrum, and the conditions probe overlapping but "
        "distinct populations of emitters. The channel groups are therefore "
        "strongly coupled and highly redundant, with a variable number of "
        "valid channels per group — an adversarial setting for "
        "marginal-scoring DR and a fair stress test for dependency-aware "
        "methods. We use two datasets: a biological-morphology set (192 "
        "channels, 4 classes) and a materials-chemistry set (158 channels, 3 "
        "classes). Selection is fully unsupervised; a downstream classifier is "
        "used only to measure how much decision-relevant information the "
        "retained channels preserve.", first_indent=0)
    body_para(doc,
        "Accuracy under reduction. On the biological set the full 192-channel "
        "baseline reaches 88.2% accuracy; the method reaches 95.2% with 80 "
        "channels (+7.0 points, 58% reduction) and still 89.4% with only 9 "
        "channels (95% reduction). On the materials set the 158-channel "
        "baseline reaches 79.8%; the method reaches 85.6% with 30 channels "
        "(+5.8 points, 81% reduction). As Fig. 3 shows, accuracy peaks at "
        "intermediate sizes: beyond the peak, additional channels reintroduce "
        "noise and degrade performance — selection acts as denoising, not "
        "merely compression. Table I summarizes the operating points.")
    column_figure(doc, FIG_ENVELOPE,
        "Fig. 3.  Accuracy versus number of retained channels (biological "
        "benchmark). The selection (best, dashed) exceeds the full-channel "
        "baseline (red) across a wide range and peaks well below full "
        "dimensionality, confirming that conditional selection denoises rather "
        "than merely compresses.")
    results_table(doc)
    body_para(doc,
        "Is the selection load-bearing? Against 10,000 random channel subsets "
        "of equal size, the learned selection scores 90.2% versus a random "
        "mean of 46.1% (max 57.9%), a separation of about 12 standard "
        "deviations. The retained channels are thus far from arbitrary.")
    body_para(doc,
        "Does it generalize? The identical pipeline, with no re-tuning, "
        "succeeds on both datasets despite their different physics. On the "
        "materials set, nine of ten downstream classifiers (k-NN, LDA, linear "
        "and RBF SVM, MLP, random forests, gradient boosting) benefit from the "
        "selection, indicating the retained information is classifier-agnostic "
        "rather than tuned to one model.")

    # ---- V. Conclusion ----
    section_heading(doc, "V", "Conclusion")
    body_para(doc,
        "We argued that effective dimensionality reduction for complex sensor "
        "data must model conditional, not marginal, channel relevance, and we "
        "introduced an unsupervised framework that does so by learning the "
        "joint channel structure, attributing conditional importance through "
        "sensitivity analysis, and selecting a minimal non-redundant subset. "
        "On a deliberately coupled benchmark it reduced channels by 80–95% "
        "while matching or exceeding full-data accuracy, with selections that "
        "are statistically far from chance and robust across classifiers. The "
        "framework is modality-agnostic; future work will apply it to "
        "multimodal IoT and industrial sensor streams, automate the choice of "
        "target size, and tighten the link between the attribution score and "
        "conditional mutual information.", first_indent=0)

    # ---- Acknowledgements (with mandatory AI disclosure) ----
    section_heading(doc, "", "Acknowledgements")
    body_para(doc,
        "The authors thank colleagues at the American University of Armenia "
        "for discussions and access to imaging data. In accordance with the "
        "CODASSCA 2026 policy on AI-generated content, the authors disclose "
        "that a large-language-model assistant (Anthropic Claude) was used for "
        "drafting and language editing of this manuscript; all research design, "
        "experiments, analyses, and results are the authors’ own and were "
        "verified by the authors.", first_indent=0)

    # ---- References ----
    section_heading(doc, "", "References")
    refs = [
        "H. Peng, F. Long, and C. Ding, “Feature selection based on "
        "mutual information: criteria of max-dependency, max-relevance, and "
        "min-redundancy,” IEEE Trans. Pattern Anal. Mach. Intell., vol. "
        "27, no. 8, pp. 1226–1238, 2005.",
        "G. E. Hinton and R. R. Salakhutdinov, “Reducing the "
        "dimensionality of data with neural networks,” Science, vol. 313, "
        "no. 5786, pp. 504–507, 2006.",
        "L. McInnes, J. Healy, and J. Melville, “UMAP: Uniform Manifold "
        "Approximation and Projection for dimension reduction,” arXiv:"
        "1802.03426, 2018.",
        "Y. Cai, X. Liu, and Z. Cai, “BS-Nets: An end-to-end framework "
        "for band selection of hyperspectral image,” IEEE Trans. Geosci. "
        "Remote Sens., vol. 58, no. 3, pp. 1969–1984, 2020.",
        "J. Carbonell and J. Goldstein, “The use of MMR, diversity-based "
        "reranking for reordering documents and producing summaries,” in "
        "Proc. ACM SIGIR, 1998, pp. 335–336.",
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        tighten(p, after=1)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        run = p.add_run(f"[{i}] {r}")
        run.font.name = BODY_FONT
        run.font.size = Pt(8.5)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
