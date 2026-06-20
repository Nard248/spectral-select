"""Generate a Word file summarising the Communications AI & Computing
submission guidelines we found and how the manuscript addresses each.

Output: CommsAIComputing_Guidelines_Compliance.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parent / "CommsAIComputing_Guidelines_Compliance.docx"

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
AMBER = RGBColor(0xB7, 0x6E, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_cell(cell, text, *, bold=False, size=9, color=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color


def h(doc, text, size=14, color=NAVY, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def body(doc, text, *, size=10, italic=False, color=None, bullet=False):
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def build():
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.7)
        s.top_margin = s.bottom_margin = Inches(0.7)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Communications AI & Computing — Submission Guidelines & Compliance")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Manuscript: “Unsupervised Deep Learning for Wavelength Selection "
                     "in Multi-Excitation Fluorescence Imaging”  ·  primary research Article")
    rs.italic = True
    rs.font.size = Pt(10)
    rs.font.color.rgb = GREY
    rs.font.name = "Calibri"

    # Sourcing caveat
    h(doc, "How this information was obtained", size=12)
    body(doc,
         "The journal’s Guide-to-Authors and Submission-Guidelines pages redirect to a "
         "Nature login wall and could not be read directly. The requirements below were "
         "recovered from cached web-search snippets of the journal’s “Content types” "
         "and “Submission guidelines” pages. They should be treated as a faithful lead, "
         "and confirmed against the live pages (which you can access when logged in) before "
         "final submission. Numbers are quoted as found.")

    # Legend
    h(doc, "Compliance summary", size=12)
    leg = doc.add_paragraph()
    for txt, col in [("✓ met   ", GREEN), ("⚠ to confirm / deviation   ", AMBER),
                     ("ℹ informational", GREY)]:
        rr = leg.add_run(txt)
        rr.font.size = Pt(9)
        rr.font.color.rgb = col
        rr.bold = True
        rr.font.name = "Calibri"

    # Table
    rows = [
        # (Requirement, What the guideline says, How the manuscript addresses it, status, statuscolor)
        ("Journal / article type",
         "Communications AI & Computing (Nature Portfolio). Primary research “Article”.",
         "Prepared as a primary research Article.", "ℹ", GREY),

        ("Main-text word limit",
         "Articles “limited to fewer than ~5,000 words of main text; the Methods section "
         "length is not limited and does not count toward the word count.”",
         "Main text (Introduction + Related Work + Results + Discussion + Conclusion) = "
         "~4,308 words. Methods (Methodology + Experimental Setup) = ~2,314 words, excluded.",
         "✓", GREEN),

        ("Page limit",
         "No page limit stated for Articles (page limits are given only for Reviews, ≤8 pp, "
         "and Comments, 1–4 pp).",
         "Not applicable. Layout is two-column, ~16 pp; governed by the word limit, which is met.",
         "✓", GREEN),

        ("Title",
         "≤15 words; should not contain technical terms, abbreviations, punctuation, or "
         "active verbs.",
         "10 words, no abbreviations, no active verb. (Note: contains the hyphenated compound "
         "“Multi-Excitation”; hyphens in compound modifiers are normally accepted.)",
         "✓", GREEN),

        ("Abstract",
         "Roughly 150–200 words; contains no references.",
         "196 words; no references.",
         "✓", GREEN),

        ("Reference style",
         "Standard Nature referencing style; list all authors unless there are six or more, "
         "in which case only the first author is given, followed by ‘et al.’",
         "54 references cited. Currently formatted with IEEEtran style (numeric). To be switched "
         "to the Nature numbered/superscript style; author-list ‘et al.’ rule to be applied.",
         "⚠", AMBER),

        ("Reference-count cap",
         "No maximum number of references found for Articles.",
         "54 references — no action needed unless the live guidelines state a cap.",
         "ℹ", GREY),

        ("Display items (figures/tables)",
         "A cap of “up to 10 display items” is stated for Reviews; no explicit cap was found "
         "for Articles.",
         "~14 figures + 7 tables (~21 display items). Within the original paper’s set. "
         "If the live Article guidelines impose a cap, some tables can be merged or moved to "
         "Supplementary.",
         "⚠", AMBER),

        ("Manuscript structure / order",
         "“Manuscripts … do not need to adhere to formatting requirements at the point of "
         "initial submission; formatting requirements only apply at the time of acceptance.” "
         "Nature house style places Methods at the end.",
         "Follows the original paper order (Introduction, Related Work, Methodology, Experimental "
         "Setup, Results, Discussion, Conclusion). Methods are NOT at the end — a deliberate "
         "deviation, permissible at initial submission, to be reformatted at acceptance.",
         "⚠", AMBER),

        ("File format",
         "A single PDF or Word file is encouraged; figures may be inline or grouped at the end; "
         "each figure legend on the same page as its figure.",
         "LaTeX source compiles to a single self-contained PDF; figures are inline with legends "
         "on the same page.",
         "✓", GREEN),

        ("AI-use disclosure",
         "Use of AI-generated content must be disclosed.",
         "Disclosed in the Acknowledgment: an LLM assistant was used only for language/source "
         "editing, not for conceptual design, analysis, or intellectual content.",
         "✓", GREEN),

        ("Data availability",
         "A data-availability statement is required.",
         "Included (“available from the corresponding author on reasonable request”). Replace "
         "with a repository DOI if the cubes can be deposited.",
         "✓", GREEN),

        ("Code availability",
         "A code-availability statement is required.",
         "Included, pointing to the GitHub repository.",
         "✓", GREEN),

        ("Author contributions",
         "A contributions statement is required.",
         "Included (drafted) — confirm the wording reflects each author’s actual role.",
         "⚠", AMBER),

        ("Competing interests",
         "A competing-interests declaration is required.",
         "Included (“The authors declare no competing interests”).",
         "✓", GREEN),

        ("Author list / affiliations",
         "Full author names and affiliations; corresponding author and ORCIDs at submission.",
         "Two authors with three affiliations and emails. Add ORCIDs and mark the corresponding "
         "author in the portal.",
         "⚠", AMBER),
    ]

    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [Inches(1.45), Inches(2.7), Inches(2.7), Inches(0.5)]
    hdr = ["Requirement", "What the guideline says", "How the manuscript addresses it", "St."]
    for j, (c, w) in enumerate(zip(table.rows[0].cells, widths)):
        set_cell(c, hdr[j], bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(c, "1F4E79")
        c.width = w
    for i, (req, says, how, st, col) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        set_cell(cells[0], req, bold=True, size=9)
        set_cell(cells[1], says, size=9)
        set_cell(cells[2], how, size=9)
        set_cell(cells[3], st, bold=True, size=11, color=col)
        for j, w in enumerate(widths):
            cells[j].width = w
        if i % 2 == 0:
            for c in cells:
                shade(c, "F2F5FA")

    # Outstanding items
    h(doc, "Items to resolve before final submission", size=12)
    for t in [
        "Switch the bibliography from IEEEtran to the Nature numbered/superscript style and "
        "apply the ‘six-or-more authors → et al.’ rule.",
        "Confirm the Author-Contributions wording and add ORCIDs / mark the corresponding author.",
        "Confirm against the live (logged-in) Guide-to-Authors whether Articles have a hard cap "
        "on display items or references; if so, consolidate tables or move some to Supplementary.",
        "The current two-column IEEE layout and mid-paper Methods are acceptable at initial "
        "submission but must be reformatted to the Nature template (single column, Methods at end) "
        "if the paper is accepted.",
        "Optionally replace the data-availability statement with a repository DOI.",
    ]:
        body(doc, t, bullet=True, size=10)

    # Footer note
    n = doc.add_paragraph()
    rn = n.add_run("Prepared as an internal checklist; not part of the submitted manuscript. "
                   "Figures-and-tables counts are approximate and include schematic (TikZ) figures.")
    rn.italic = True
    rn.font.size = Pt(8)
    rn.font.color.rgb = GREY
    rn.font.name = "Calibri"

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
