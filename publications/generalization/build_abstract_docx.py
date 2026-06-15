"""Generate the submittable Word document for the general channel-selection paper.

Framing (per author direction):
- One general method, verified in TWO domains: wearable-sensor HAR (the new, detailed
  result) and biomedical hyperspectral imaging (brief, existing verified results).
- Biomedical is NOT framed as prior work; it is a second verification domain.
- Plain language, simple visuals, explicit selection -> verification flow, dataset detail.
- Honest baselines: matches supervised / beats variance / more stable; no "beats random",
  no "first" claims. Biomedical numbers are existing verified results.

Run: python generalization/build_abstract_docx.py
Output: generalization/Meloyan_GeneralChannelSelection.docx
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
FIGS = HERE / "figures"
OUT = HERE / "Meloyan_GeneralChannelSelection.docx"
BODY_FONT = "Times New Roman"
BODY_PT = 9.5


def set_columns(section, num, space_twips=360):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols"); sectPr.append(cols)
    cols.set(qn("w:num"), str(num)); cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")


def tighten(p, before=0, after=2, line=1.0):
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line


def _margins(section):
    section.top_margin = Inches(0.75); section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.625); section.right_margin = Inches(0.625)


def body_para(doc, text, *, justify=True, first_indent=0.18, size=BODY_PT, bold=False, italic=False):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    tighten(p)
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(first_indent)
    r = p.add_run(text)
    r.font.name = BODY_FONT; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p


def section_heading(doc, numeral, title):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(p, before=6, after=2)
    label = f"{numeral}. {title.upper()}" if numeral else title.upper()
    r = p.add_run(label)
    r.font.name = BODY_FONT; r.font.size = Pt(BODY_PT); r.bold = True; r.font.small_caps = True


def _caption(doc, text):
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(cap, before=1, after=4)
    r = cap.add_run(text); r.font.name = BODY_FONT; r.font.size = Pt(8)


def column_figure(doc, img, caption, width_in=3.35):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(p, before=3, after=0)
    if Path(img).exists():
        p.add_run().add_picture(str(img), width=Inches(width_in))
    else:
        p.add_run(f"[missing figure: {Path(img).name}]")
    _caption(doc, caption)


def _begin_fullwidth(doc):
    s = doc.add_section(WD_SECTION.CONTINUOUS); _margins(s); set_columns(s, 1)


def _end_fullwidth(doc):
    s = doc.add_section(WD_SECTION.CONTINUOUS); _margins(s); set_columns(s, 2)


def fullwidth_figure(doc, img, caption, width_in=6.9):
    _begin_fullwidth(doc)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(p, before=2, after=0)
    if Path(img).exists():
        p.add_run().add_picture(str(img), width=Inches(width_in))
    else:
        p.add_run(f"[missing figure: {Path(img).name}]")
    _caption(doc, caption)
    _end_fullwidth(doc)


def table_block(doc, caption, rows, col_w=None, fullwidth=True):
    if fullwidth:
        _begin_fullwidth(doc)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tighten(cap, before=4, after=1)
    rc = cap.add_run(caption); rc.font.name = BODY_FONT; rc.font.size = Pt(8); rc.bold = True
    rc.font.small_caps = True
    table = doc.add_table(rows=len(rows), cols=len(rows[0])); table.alignment = 1
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]; cell.text = ""
            par = cell.paragraphs[0]; par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tighten(par, before=0, after=0)
            run = par.add_run(str(val)); run.font.name = BODY_FONT; run.font.size = Pt(8)
            run.bold = (i == 0)
    if fullwidth:
        _end_fullwidth(doc)
    return table


ABSTRACT = (
    "We present a method that picks a small, useful set of sensor channels without using any "
    "labels. A group-structured autoencoder first learns a shared representation of all the "
    "channels. We then perturb its latent factors and watch how each channel's reconstruction "
    "responds. A channel that responds strongly is one the model relies on, so this gives a "
    "label-free measure of relevance. A relevance and redundancy rule turns these scores into an "
    "ordered shortlist of the real channels, avoiding picks that merely repeat information "
    "already kept. The same method applies to very different sensors, and only the front-end "
    "encoder changes between them. On a wearable-sensor activity benchmark with three inertial "
    "units, keeping ten of twenty-seven channels holds accuracy to within 0.04 macro-F1 of the "
    "full set. It also equals a supervised selector that does use labels, and it is far more "
    "stable across people than variance ranking. On biomedical hyperspectral imaging the same "
    "method removes 58 to 95 percent of channels while holding or improving accuracy. Accuracy "
    "rises from 88.2 to 95.2 percent on lichen tissue and from 79.8 to 85.6 percent on collagen. "
    "One simple, label-free idea therefore reduces channels well across markedly different "
    "sensing modalities."
)

REFS = [
    "M. F. Balin, A. Abid, and J. Zou, \"Concrete autoencoders: Differentiable feature "
    "selection and reconstruction,\" in Proc. ICML, 2019.",
    "H. Peng, F. Long, and C. Ding, \"Feature selection based on mutual information,\" IEEE "
    "Trans. Pattern Anal. Mach. Intell., vol. 27, no. 8, 2005.",
    "X. He, D. Cai, and P. Niyogi, \"Laplacian score for feature selection,\" in Proc. "
    "NeurIPS, 2005.",
    "Y. Cai et al., \"BS-Nets: An end-to-end framework for band selection of hyperspectral "
    "image,\" IEEE Trans. Geosci. Remote Sens., vol. 58, no. 3, 2020.",
    "A. Reiss and D. Stricker, \"Introducing a new benchmarked dataset for activity "
    "monitoring,\" in Proc. IEEE Int. Symp. Wearable Computers (ISWC), 2012.",
    "A. Carbonneau et al., \"The MONSTER benchmark for multivariate time-series "
    "classification,\" arXiv:2502.15122, 2025.",
]


def build():
    doc = Document()
    sec = doc.sections[0]; sec.page_height = Inches(11); sec.page_width = Inches(8.5); _margins(sec)

    # Title
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER; tighten(t, after=4)
    rt = t.add_run("Label-Free, Dependency-Aware Channel Selection\n"
                   "for Coupled Multi-Channel Sensor Data")
    rt.font.name = BODY_FONT; rt.font.size = Pt(18); rt.bold = True
    a = doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER; tighten(a, before=2, after=0)
    ra = a.add_run("Narek Meloyan, Aleksandr Hayrapetyan, and Narine Sarvazyan")
    ra.font.name = BODY_FONT; ra.font.size = Pt(11)
    af = doc.add_paragraph(); af.alignment = WD_ALIGN_PARAGRAPH.CENTER; tighten(af, before=1, after=8)
    raf = af.add_run("American University of Armenia, Yerevan, Armenia\n"
                     "{narek.meloyan, aleksandr.hayrapetyan, nsarvazyan}@aua.am")
    raf.font.name = BODY_FONT; raf.font.size = Pt(9); raf.italic = True

    body = doc.add_section(WD_SECTION.CONTINUOUS); _margins(body); set_columns(body, 2)

    # Abstract + index terms
    ab = doc.add_paragraph(); ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; tighten(ab, after=3)
    lab = ab.add_run("Abstract—"); lab.font.name = BODY_FONT; lab.font.size = Pt(BODY_PT)
    lab.bold = True; lab.italic = True
    rab = ab.add_run(ABSTRACT); rab.font.name = BODY_FONT; rab.font.size = Pt(BODY_PT)
    rab.bold = True; rab.italic = True
    idx = doc.add_paragraph(); idx.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; tighten(idx, after=4)
    il = idx.add_run("Index Terms—"); il.font.name = BODY_FONT; il.font.size = Pt(BODY_PT)
    il.bold = True; il.italic = True
    it = idx.add_run("unsupervised feature selection, channel selection, conditional relevance, "
                     "autoencoders, interpretability, human activity recognition, hyperspectral "
                     "imaging.")
    it.font.name = BODY_FONT; it.font.size = Pt(BODY_PT); it.italic = True

    # I. Introduction
    section_heading(doc, "I", "Introduction")
    body_para(doc,
        "Complex sensors emit many channels that overlap heavily and arrive in natural groups, "
        "such as the three axes of one inertial unit or the emission bands measured under one "
        "excitation. "
        "For cost, power, transmission, and interpretability, one often wants to keep only a "
        "handful of the actual channels, not a learned mixture of them. The challenge is to decide "
        "which channels to keep without labels.")
    body_para(doc,
        "Existing selectors force a trade-off. Unsupervised filters such as variance ranking or "
        "PCA score each channel by its own statistics and ignore how channels depend on one "
        "another, so they tend to keep several channels that carry the same information. "
        "Dependency-aware selectors such as mRMR or JMI account for this redundancy but need "
        "labels. The cell that is at once label-free, dependency-aware, and discrete (returning "
        "real channels, not a projection) is sparsely occupied. This paper presents a method for "
        "that cell and, more importantly, shows that one such method works across two very "
        "different sensing modalities by changing only its front-end encoder.")

    # II. Method
    section_heading(doc, "II", "Method")
    body_para(doc,
        "The method has three simple stages (Fig. 1). (1) A group-structured autoencoder is "
        "trained by reconstruction only, with no labels. Each channel group has its own small "
        "encoder, and the per-group features are averaged into one shared latent representation, "
        "which is then decoded back to every group. (2) We probe what the model learned. Nudging "
        "one latent factor and measuring how much each channel's reconstruction changes gives that "
        "channel's relevance, and channels the model depends on react strongly. (3) A "
        "relevance-redundancy rule "
        "(maximal marginal relevance) then picks channels one at a time, preferring those that are "
        "informative yet not already represented by a previously chosen channel. The output is an "
        "ordered, interpretable list of real channels that can be cut to any target size.")
    fullwidth_figure(doc, FIGS / "fig_method.png",
        "Fig. 1.  The three-stage pipeline. The selection engine (perturbation -> per-channel "
        "relevance -> relevance-redundancy selection) is identical across domains; only the "
        "encoder/decoder changes to match the data's regular axis.")
    body_para(doc,
        "Crucially, the selection engine is modality-agnostic: it works on an abstract layout of "
        "groups, channels, and one regular axis (space or time). Moving to a new modality requires "
        "only an encoder matched to that axis (Fig. 2). We verify the same engine on two domains.")
    fullwidth_figure(doc, FIGS / "fig_crossdomain.png",
        "Fig. 2.  One method, two verification domains. Only the encoder changes (1-D temporal "
        "convolutions for wearable sensors; 2-D/3-D spatial convolutions for imaging).")

    # III. Verification 1 — HAR (detailed)
    section_heading(doc, "III", "Verification 1: Wearable-Sensor Activity Recognition")
    body_para(doc,
        "Dataset. We use PAMAP2, a standard human-activity-recognition (HAR) benchmark recorded "
        "from three body-worn inertial measurement units (IMUs) placed on the hand, chest, and "
        "ankle, sampled at 100 Hz. Each IMU contributes nine channels (3-axis accelerometer, "
        "gyroscope, and magnetometer), giving 27 channels organized into three natural groups. We "
        "use the preprocessed release with fixed one-second windows, yielding 38,856 windows "
        "across eight subjects and twelve activity classes.")
    body_para(doc,
        "Protocol (the selection-then-verification flow). For each held-out subject we (i) train "
        "the group autoencoder unsupervised on the remaining subjects' windows, with one IMU per "
        "group and a 1-D temporal-convolution encoder; (ii) run the selection engine to obtain an "
        "ordered channel list; and (iii) verify the chosen subset by training a k-nearest-neighbour "
        "classifier on those channels and measuring macro-F1 on the held-out subject. All numbers "
        "use leave-one-subject-out (LOSO) cross-validation, the rigorous protocol in which the "
        "test subject is never seen during training. We compare against a supervised "
        "mutual-information selector, "
        "an unsupervised variance baseline, and random selection.")
    column_figure(doc, FIGS / "fig_acc_vs_k.png",
        "Fig. 3.  Macro-F1 versus number of retained channels (LOSO). The label-free method (red) "
        "tracks the supervised selector (blue) and stays above variance (orange).")
    body_para(doc,
        "Results (Table I). Reducing from 27 to 10 channels (a 63% reduction) keeps macro-F1 within "
        "0.04 of the full set, and the label-free method matches the supervised selector at every "
        "budget while clearly beating variance. It also selects far more consistently across "
        "subjects (Fig. 4). Variance swings widely from person to person, whereas our method does "
        "not, which matters when a single sensor set must be fixed before deployment. For full "
        "transparency, on this particular benchmark no selector, including the supervised one, "
        "surpasses random selection, because the 27 channels are highly redundant and almost any "
        "moderate subset approaches the full-set ceiling; the informative comparisons here are "
        "therefore against the supervised selector and on stability.")
    table_block(doc, "TABLE I.  HAR macro-F1 under channel reduction (LOSO). Full set (27 ch) = 0.72.",
        [("Method", "10 ch (-63%)", "7 ch (-74%)", "5 ch (-81%)"),
         ("Ours (label-free)", "0.68", "0.63", "0.54"),
         ("Mutual info (supervised)", "0.68", "0.61", "0.54"),
         ("Variance (unsupervised)", "0.59", "0.54", "0.51")])
    column_figure(doc, FIGS / "fig_stability.png",
        "Fig. 4.  Stability of the selected set across subjects (lower is better). The proposed "
        "method is markedly more consistent than variance selection.")

    # IV. Verification 2 — Biomedical HSI (brief)
    section_heading(doc, "IV", "Verification 2: Biomedical Hyperspectral Imaging")
    body_para(doc,
        "We verify the identical method on biomedical multi-excitation fluorescence hyperspectral "
        "imaging, where a group is an excitation wavelength, a channel is an emission band, and the "
        "regular axis is the 2-D image plane (a spatial-convolution encoder is substituted). On a "
        "lichen-tissue dataset (192 bands, four classes) the method reaches 95.2% classification "
        "accuracy using 80 of 192 bands, above the 88.2% obtained with all bands, and still "
        "achieves 89.4% using only nine bands (a 95% reduction). On a collagen dataset (158 bands, "
        "three "
        "classes), 30 bands raise accuracy from 79.8% to 85.6% (Table II, Fig. 5). In this "
        "high-redundancy imaging regime, intelligent selection clearly separates from trivial "
        "baselines, complementing the wearable-sensor result.")
    table_block(doc, "TABLE II.  Biomedical HSI: aggressive reduction maintains or improves accuracy.",
        [("Dataset", "Channels (all -> kept)", "Reduction", "Accuracy (all -> kept)"),
         ("Lichens", "192 -> 80", "58%", "88.2% -> 95.2%"),
         ("Lichens", "192 -> 9", "95%", "88.2% -> 89.4%"),
         ("Collagen", "158 -> 30", "81%", "79.8% -> 85.6%")])
    column_figure(doc, FIGS / "fig_biomed.png",
        "Fig. 5.  Biomedical HSI: fewer channels, same-or-better accuracy.")

    # V. Conclusion
    section_heading(doc, "V", "Conclusion and Future Work")
    body_para(doc,
        "A single unsupervised, dependency-aware selection principle delivers strong, interpretable "
        "channel reduction across two very different modalities, requiring only an encoder change. "
        "On wearable sensors it matches a supervised selector without labels and selects more "
        "stably than variance; on biomedical imaging it cuts 58-95% of channels while maintaining "
        "or improving accuracy. Future work targets a higher-redundancy activity benchmark, where "
        "selection should separate from random by a wide margin, and a direct comparison with "
        "differentiable unsupervised selectors such as the concrete autoencoder.")

    # VI. Acknowledgement
    section_heading(doc, "VI", "Acknowledgement")
    body_para(doc,
        "AI-assisted tooling was used for code scaffolding, figure generation, and manuscript "
        "drafting; all experimental design, analysis, and claims were verified by the authors.")

    # References
    section_heading(doc, "", "References")
    for i, r in enumerate(REFS, 1):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; tighten(p, after=1)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        run = p.add_run(f"[{i}] {r}"); run.font.name = BODY_FONT; run.font.size = Pt(8)

    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    build()
