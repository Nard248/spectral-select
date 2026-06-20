"""Build the GUI User Guide .docx from content.json + the rendered screenshots.

Run:  python docs/gui_user_guide/build_docx.py
Output: docs/gui_user_guide/spectral-select_GUI_User_Guide.docx
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).parent
LIVE = HERE / "screenshots_live"      # real-data run (preferred)
INITIAL = HERE / "screenshots"        # empty-state fallback
CONTENT = HERE / "content.json"
OUT = HERE / "spectral-select_GUI_User_Guide.docx"


def shot_path(name: str) -> Path:
    """Prefer the live (real-data) screenshot; fall back to the initial-state one."""
    live = LIVE / name
    return live if live.exists() else (INITIAL / name)

SLUGS = {
    1: "load", 2: "metadata", 3: "normalize", 4: "spatial_crop", 5: "spectral_crop",
    6: "draw_classes", 7: "roi_regions", 8: "export", 9: "train", 10: "select",
}

# Corrected navigation note for Step 8 (the auto-extracted one mislabeled steps 2-7).
NAV_OVERRIDES = {
    8: ("This is the final preprocessing step. Export is enabled once you have loaded data "
        "(Step 1) and is most useful after you have defined a class mask (Step 6) and ROI "
        "regions (Step 7). After a successful export, the saved spectra feed directly into "
        "Step 9 (Train Autoencoder) — or can be shared with collaborators. Steps 9 and 10 "
        "(train + select) follow Export in this combined workflow."),
}


def _clean(text: str) -> str:
    return (text or "").strip().rstrip('"').rstrip("\\").strip()


def add_picture(doc: Document, path: Path, width_in: float = 6.5) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[screenshot missing: {path.name}]")


def controls_table(doc: Document, controls: list[dict]) -> None:
    table = doc.add_table(rows=1, cols=2)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Control"
    hdr[1].text = "What it does"
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for c in controls:
        row = table.add_row().cells
        row[0].text = c.get("control", "")
        row[1].text = c.get("does", "")
    # widen the description column
    for row in table.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.3)


def main() -> None:
    steps = json.loads(CONTENT.read_text())
    doc = Document()

    # ---- Title page ----
    title = doc.add_heading("spectral-select", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("GUI User Guide — MEHSI Preprocessor")
    run.bold = True
    run.font.size = Pt(18)
    tagline = doc.add_paragraph(
        "A no-code, step-by-step workflow for preparing 4D hyperspectral data and running "
        "Perturbation-Based Autoencoder wavelength selection."
    )
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_picture(doc, shot_path("overview.png"), width_in=6.3)
    cap = doc.add_paragraph("The MEHSI Preprocessor wizard — a 10-step sidebar on the left, "
                            "the active page on the right.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    # ---- About ----
    doc.add_heading("About this tool", level=1)
    doc.add_paragraph(
        "The MEHSI Preprocessor is a desktop application that takes you from raw multi-excitation "
        "hyperspectral image files all the way to a ranked set of the most informative wavelength "
        "bands — without writing any code. The first eight steps load and clean your data and let "
        "you annotate it; the final two steps train a convolutional autoencoder and then run the "
        "Perturbation-Based Autoencoder (AE) selection that identifies which wavelengths carry the "
        "most information."
    )
    doc.add_paragraph(
        "How the selection works, in one sentence: the autoencoder learns a compressed "
        "representation of your spectra, the method nudges (perturbs) that representation and "
        "measures how much each wavelength band changes in response, and the most influential, "
        "least-redundant bands are selected.",
    )

    doc.add_heading("Worked example in this guide", level=2)
    doc.add_paragraph(
        "Every screenshot below comes from a real, end-to-end run on the Collagen (acetic acid) "
        "dataset — 6 excitations (310–400 nm, 31 bands each). The data was loaded from raw .im3 "
        "files, normalized by exposure and laser power, spatially cropped to an 80×80 region, "
        "spectrally filtered with a Rayleigh cutoff (158 bands kept), annotated with three classes, "
        "exported, used to train a convolutional autoencoder, and finally run through band selection "
        "to rank the 12 most informative wavelengths. So what you see on each page is the tool "
        "actually working on real data, not mock-ups."
    )

    # ---- Install & launch ----
    doc.add_heading("Installation & launch", level=1)
    doc.add_paragraph("Install the package with its GUI extra, then launch the app:")
    code = doc.add_paragraph()
    code.add_run('pip install -e ".[gui]"\nspectral-select-gui      # or:  python -m mehsi_preprocessor').font.name = "Courier New"
    doc.add_paragraph(
        "The window opens on Step 1. Work top-to-bottom through the sidebar."
    )

    # ---- Navigation ----
    doc.add_heading("Navigating the wizard", level=1)
    for line in [
        "The numbered list on the left (“Steps”) is your map. Click any step to jump to it, "
        "or use the Previous / Next buttons at the bottom-right.",
        "Steps build on each other. Later steps stay locked or empty until the steps they depend on "
        "are done — for example, Select Bands (Step 10) is disabled until you train or load a model "
        "in Step 9.",
        "Editing an earlier step invalidates the results of later steps. If you re-crop or re-train, "
        "you will need to re-run the steps after it. This keeps your results consistent with your "
        "current data.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("The ten steps", level=2)
    for s in steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{s['title']} — ").bold = True
        p.add_run(s["purpose"].split(". ")[0].strip() + ".")
    doc.add_page_break()

    # ---- Per-step sections ----
    for s in steps:
        n = s["step_number"]
        doc.add_heading(f"Step {n} — {s['title']}", level=1)

        doc.add_heading("Purpose", level=3)
        doc.add_paragraph(s["purpose"])

        add_picture(doc, shot_path(f"step_{n:02d}_{SLUGS.get(n, 'page')}.png"))
        cap = doc.add_paragraph(
            f"Step {n}: {s['title']} — shown with the Collagen (acetic acid) sample loaded."
        )
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_heading("What to do", level=3)
        for action in s["what_to_do"]:
            doc.add_paragraph(action, style="List Number")

        doc.add_heading("Controls on this page", level=3)
        controls_table(doc, s.get("key_controls", []))

        doc.add_heading("Navigation", level=3)
        doc.add_paragraph(NAV_OVERRIDES.get(n, _clean(s.get("navigation_note", ""))))

        pitfalls = s.get("pitfalls") or []
        if pitfalls:
            doc.add_heading("Tips & pitfalls", level=3)
            for tip in pitfalls:
                doc.add_paragraph(tip, style="List Bullet")

        doc.add_page_break()

    # ---- Quick recipe ----
    doc.add_heading("End-to-end quick recipe", level=1)
    recipe = [
        "Step 1 — Browse to your folder of .im3 files and let them load.",
        "Step 2 — Confirm exposure times and laser powers (or browse to the metadata files).",
        "Step 3 — Click Apply Normalization.",
        "Step 4 — Draw a rectangle to keep, then Apply Crop.",
        "Step 5 — Set the Rayleigh cutoff / emission ranges, Preview, then Apply Spectral Crop.",
        "Step 6 — Add classes and paint them onto the image.",
        "Step 7 — Draw ROI rectangles and assign each to a class.",
        "Step 8 — Choose an output folder and Export.",
        "Step 9 — Train a new autoencoder (or load a .pth) and wait for “Model ready ✓”.",
        "Step 10 — Set the number of bands, Run selection, review the table, and Export results.",
    ]
    for r in recipe:
        doc.add_paragraph(r, style="List Number")

    doc.add_heading("Where your files go", level=2)
    doc.add_paragraph(
        "Step 8 writes spectra_masked.pkl, spectra_unmasked.pkl, class_mask.png and "
        "roi_regions.json to the output folder you choose. Step 10 writes the selected-band "
        "results (CSV / JSON / TIFF) to the folder you choose when you click Export results."
    )

    doc.save(str(OUT))
    print(f"Saved {OUT}  ({OUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
