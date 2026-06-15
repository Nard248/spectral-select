#!/usr/bin/env python3
"""
Build IASIM 2026 Abstract (Pepsin Collagen) from Official Template
====================================================================
Uses the Downloads/IASIM2026-abstract-template.dotx as the starting point,
replaces placeholder content with our Pepsin collagen study, and embeds the
accuracy envelope figure as Fig. 1.

Output: results/Pepsin_Paper_Figures/IASIM2026_Pepsin_Abstract.docx

Content decisions (see comments in build_abstract() for the reasoning):
  - Figure choice: accuracy envelope ('coverage of accuracies with peak')
  - Word count target: ~220 words body text (matches IASIM conventions)
  - Title: emphasizes the method + the validation domain (collagen)
  - Authors/affiliations: left editable so the user can finalize
"""

import zipfile
import shutil
from pathlib import Path
import warnings
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

warnings.filterwarnings('ignore')

PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_DOTX = Path("/Users/narekmeloyan/Downloads/IASIM2026-abstract-template.dotx")
REFERENCE_DOCX = Path("/Users/narekmeloyan/Downloads/Villarruel abstract NAS.docx")

OUT_DIR = PROJECT_ROOT / "results" / "Pepsin_Paper_Figures"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_DOCX = OUT_DIR / "IASIM2026_Pepsin_Abstract.docx"
FIGURE_PNG = OUT_DIR / "iasim_figure.png"

PIPELINE_CSV = PROJECT_ROOT / "results" / "Collagen_Pepsin_Normalized" / "results.csv"


# ═══════════════════════════════════════════════════════════════════════════
# FIGURE PREPARATION
# ═══════════════════════════════════════════════════════════════════════════

def build_figure():
    """Accuracy envelope with peak annotation — the 'coverage + peak' plot."""
    df = pd.read_csv(PIPELINE_CSV)
    bl = df[df['config'] == 'BASELINE'].iloc[0]
    sel = df[df['config'] != 'BASELINE']
    stats = sel.groupby('n_features').agg(
        mn=('accuracy', 'min'), mx=('accuracy', 'max'), mean=('accuracy', 'mean')
    ).reset_index().sort_values('n_features')

    # Compact aspect ratio suitable for a single-column abstract figure
    fig, ax = plt.subplots(figsize=(6.5, 3.6))

    ax.fill_between(stats['n_features'], stats['mn'], stats['mx'],
                    alpha=0.25, color='#4A90D9', label='Range (min\u2013max)')
    ax.plot(stats['n_features'], stats['mean'],
            color='black', lw=1.8, label='Mean')
    ax.plot(stats['n_features'], stats['mx'],
            color='#1B7A2B', lw=2.2, linestyle='--', label='Best')
    ax.axhline(y=bl['accuracy'], color='red', linestyle='--', lw=1.3,
               label=f'Baseline ({bl["accuracy"]:.1%})')

    # Mark peak with a star and a label placed in empty space
    # (far below the curves, so it cannot overlap the data band).
    bi = stats['mx'].idxmax()
    bn, ba = stats.loc[bi, 'n_features'], stats.loc[bi, 'mx']
    ax.scatter([bn], [ba], color='gold', s=180, zorder=6, marker='*',
               edgecolor='#1B7A2B', lw=1.4)
    ax.annotate(f"Peak {ba:.1%}\n({int(bn)} bands)",
                xy=(bn, ba),
                xytext=(85, 0.56),          # empty lower-right region
                fontsize=10, fontweight='bold', color='#1B7A2B',
                ha='center', va='center',
                arrowprops=dict(arrowstyle='-|>', color='#1B7A2B', lw=1.2,
                                shrinkA=0, shrinkB=6),
                bbox=dict(boxstyle='round,pad=0.3', facecolor='white',
                          edgecolor='#1B7A2B', alpha=0.95),
                zorder=7)

    ax.set_xlabel('Number of Bands Selected', fontsize=11)
    ax.set_ylabel('Classification Accuracy', fontsize=11)
    # Put the legend in the top-right — that area is empty because all
    # curves flatten near 0.82 by n=100, leaving space above them.
    ax.legend(loc='upper right', fontsize=8.5, framealpha=0.95)
    ax.set_xlim(0, 160)
    ax.set_ylim(0.45, 0.92)
    ax.grid(True, alpha=0.3)
    ax.tick_params(axis='both', which='major', labelsize=10)

    fig.tight_layout()
    fig.savefig(FIGURE_PNG, dpi=400, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    print(f"  Saved figure: {FIGURE_PNG.name}  ({FIGURE_PNG.stat().st_size/1024:.0f} KB)")


# ═══════════════════════════════════════════════════════════════════════════
# TEMPLATE LOADING
# ═══════════════════════════════════════════════════════════════════════════

def load_template_as_docx():
    """Convert the .dotx template to a working .docx copy."""
    working = OUT_DIR / "_template_working.docx"
    contents = {}
    with zipfile.ZipFile(TEMPLATE_DOTX, 'r') as zin:
        for name in zin.namelist():
            contents[name] = zin.read(name)
    # Convert template content-type to document content-type
    ct = contents['[Content_Types].xml'].decode('utf-8')
    ct = ct.replace(
        'application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml',
    )
    contents['[Content_Types].xml'] = ct.encode('utf-8')
    with zipfile.ZipFile(working, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in contents.items():
            zout.writestr(name, data)
    return working


# ═══════════════════════════════════════════════════════════════════════════
# ABSTRACT CONTENT
# ═══════════════════════════════════════════════════════════════════════════

# Title chosen to be substantively divergent from the main paper title
# ("Deep Learning for Dimensionality Reduction in Multi-Excitation
# Hyperspectral Imaging") while accurately describing what is new here:
# (1) the application — pepsin-digested collagen scaffolds,
# (2) the contribution — cross-classifier validation showing the selected
#     bands are useful regardless of which classifier is used downstream.
TITLE = (
    "Cross-Classifier Wavelength Compression of Multi-Excitation "
    "Fluorescence Imaging of Pepsin-Digested Collagen Scaffolds"
)

# Author list as a sequence of (text, is_superscript) runs so we get
# proper Word superscripts (not Unicode lookalikes).
#   Narek Meloyan         -> 1 (AUA), 2 (Orbeli)
#   Kristina Ghahramanyan -> 2 (Orbeli only)
#   Narine Sarvazyan      -> 1 (AUA), 2 (Orbeli), 3 (GWU)
AUTHOR_RUNS = [
    ("Narek Meloyan", False),
    ("1,2", True),
    (", Kristina Ghahramanyan", False),
    ("2", True),
    (", Narine Sarvazyan", False),
    ("1,2,3", True),
]

# Each affiliation starts with its superscript number as a real superscript.
AFFILIATION_RUNS = [
    [("1", True),
     ("American University of Armenia, Yerevan, Armenia", False)],
    [("2", True),
     ("L. A. Orbeli Institute of Physiology NAS RA, Yerevan, Armenia", False)],
    [("3", True),
     ("George Washington University, Washington, DC, USA", False)],
]

# ~150 words — mirrors the paper's framing while highlighting the new
# dataset (pepsin-digested collagen) as an extension of the framework.
PARAGRAPH_1 = (
    "Multi-excitation hyperspectral imaging (ME-HSI) creates a 4D dataset "
    "by varying the excitation wavelength of an emission-based HSI "
    "acquisition, which improves the discrimination of materials that "
    "contain mixed fluorophores. The resulting cubes are rich but highly "
    "redundant, and band selection methods developed for 3D HSI [1] fail "
    "to capture the nonlinear cross-excitation correlations intrinsic to "
    "ME-HSI. We previously introduced a deep-learning framework that "
    "addresses this gap in three stages: (1) a 3D convolutional autoencoder "
    "[2] with parallel excitation branches learns a unified latent "
    "representation of the 4D cube, (2) perturbation-based attribution "
    "traces reconstruction sensitivity to individual excitation\u2013emission "
    "wavelength pairs, and (3) Maximum Marginal Relevance (MMR) selection "
    "balances informativeness with spectral diversity [3]. Here we extend "
    "the framework to a new, chemically subtle sample: pepsin-digested "
    "collagen scaffolds with three crosslinker concentrations (0.002%, "
    "0.0075%, 0.013%), imaged at six excitations (310\u2013400 nm) with "
    "emissions 420\u2013720 nm (10 nm step), yielding 158 excitation\u2013"
    "emission band pairs per pixel."
)

FIGURE_CAPTION = (
    "Fig. 1. Classification accuracy (KNN, k=5) versus number of selected "
    "bands across 432 configurations on the pepsin-digested collagen dataset. "
    "Blue band = min\u2013max range; black = mean; green dashed = best per "
    "band count; red dashed = full-spectrum baseline (79.8%). Peak accuracy "
    "(85.6%) is reached at 30 bands \u2014 an 81% data reduction."
)

# ~120 words — focuses on results and validation across classifiers.
PARAGRAPH_2 = (
    "Across 432 configurations, every band count between 5 and 130 exceeded "
    "the 158-band KNN baseline of 79.8%; peak accuracy reached 85.6% with "
    "30 bands (81% compression, Fig. 1). The same selected wavelengths were "
    "evaluated with nine additional classifiers (LDA, linear and RBF SVM, "
    "random forests, gradient boosting, MLP). Linear Discriminant Analysis "
    "achieved 84.6% with only 5 bands and 92.5% with 50 bands, matching its "
    "full-spectrum baseline (92.8%) while using 68% fewer bands; all ten "
    "classifiers matched or exceeded their own baselines at 30\u201350 bands, "
    "confirming that the selected wavelengths carry classifier-independent "
    "discriminative information. These results demonstrate that the "
    "framework generalises from distinct biological classes (lichens) to "
    "subtle chemical-concentration contrasts, enabling shorter acquisitions "
    "with no loss of discriminative power."
)

REFERENCES = [
    # [1] Band-selection review (Sun & Du) — cited as [4] in the main paper.
    "1. W. Sun, Q. Du. Hyperspectral band selection: a review, IEEE Geosci. "
    "Remote Sens. Mag. 7 (2019) 118\u2013139.",
    # [2] Autoencoders for dimensionality reduction (Hinton & Salakhutdinov)
    #     — cited as [13] in the main paper.
    "2. G. E. Hinton, R. R. Salakhutdinov. Reducing the dimensionality of "
    "data with neural networks, Science 313 (2006) 504\u2013507.",
    # [3] MMR selection (Carbonell & Goldstein) — cited as [15] in the main paper.
    "3. J. Carbonell, J. Goldstein. The use of MMR, diversity-based reranking "
    "for reordering documents and producing summaries, Proc. ACM SIGIR (1998) "
    "335\u2013336.",
]


# ═══════════════════════════════════════════════════════════════════════════
# DOCX BUILDING
# ═══════════════════════════════════════════════════════════════════════════

def _set_paragraph_style(p, style_name):
    """Safely assign a style name; fall back to Normal if not found."""
    try:
        p.style = p.part.document.styles[style_name]
    except KeyError:
        pass


def build_abstract():
    print("\n[1/3] Preparing figure...")
    build_figure()

    print("\n[2/3] Loading template...")
    working = load_template_as_docx()
    doc = Document(str(working))

    # Clear existing body paragraphs (we'll rebuild from scratch using template styles)
    body = doc.element.body
    sect_pr = body.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
    for child in list(body):
        if child is sect_pr:
            continue  # keep page setup
        body.remove(child)

    print("[3/3] Inserting new content...")

    # Title
    p = doc.add_paragraph(TITLE)
    _set_paragraph_style(p, 'Presentation title')

    # Authors (with real Word superscripts for affiliation numbers)
    p = doc.add_paragraph()
    _set_paragraph_style(p, 'List of authors')
    for text, is_sup in AUTHOR_RUNS:
        r = p.add_run(text)
        if is_sup:
            r.font.superscript = True

    # Affiliations (each line begins with a superscript number)
    for runs in AFFILIATION_RUNS:
        p = doc.add_paragraph()
        _set_paragraph_style(p, 'List of authors')
        for text, is_sup in runs:
            r = p.add_run(text)
            if is_sup:
                r.font.superscript = True

    # Abstract body paragraph 1
    p = doc.add_paragraph(PARAGRAPH_1)
    _set_paragraph_style(p, 'Text of abstract')

    # Figure
    fig_p = doc.add_paragraph()
    _set_paragraph_style(fig_p, 'Figure caption')
    run = fig_p.add_run()
    # Width ~6 inches matches the body text column of the template
    run.add_picture(str(FIGURE_PNG), width=Inches(5.8))
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Figure caption
    p = doc.add_paragraph(FIGURE_CAPTION)
    _set_paragraph_style(p, 'Figure caption')

    # Abstract body paragraph 2
    p = doc.add_paragraph(PARAGRAPH_2)
    _set_paragraph_style(p, 'Text of abstract')

    # References header
    p = doc.add_paragraph("References")
    _set_paragraph_style(p, 'Text of abstract')

    # Reference list
    for ref in REFERENCES:
        p = doc.add_paragraph(ref)
        _set_paragraph_style(p, 'Reference list')

    # Move sectPr back to the end
    if sect_pr is not None:
        body.remove(sect_pr)
        body.append(sect_pr)

    doc.save(str(OUT_DOCX))
    # Clean up working template
    working.unlink()

    # Word count sanity check
    body_words = len(PARAGRAPH_1.split()) + len(PARAGRAPH_2.split())
    print(f"\n  Abstract body word count: {body_words} (target ~200-250)")
    print(f"  Paragraph 1: {len(PARAGRAPH_1.split())} words")
    print(f"  Paragraph 2: {len(PARAGRAPH_2.split())} words")
    print(f"\n  Saved: {OUT_DOCX}")
    print(f"  Figure: {FIGURE_PNG}")


if __name__ == '__main__':
    build_abstract()
